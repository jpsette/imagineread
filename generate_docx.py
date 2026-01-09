#!/usr/bin/env python3
"""
Script to convert markdown to plain DOCX
Uses only standard libraries
"""
import re
from pathlib import Path

def markdown_to_text(md_content):
    """Convert markdown to plain text with basic formatting"""
    # Remove code blocks
    text = re.sub(r'```[\s\S]*?```', '', md_content)
    
    # Convert headers (keep hierarchy with indentation)
    text = re.sub(r'^### (.*?)$', r'   \1', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*?)$', r'  \1', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.*?)$', r'\1', text, flags=re.MULTILINE)
    
    # Remove bold/italic
    text = re.sub(r'\*\*\*([^\*]+)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^\*]+)\*', r'\1', text)
    text = re.sub(r'___([^_]+)___', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Convert lists
    text = re.sub(r'^\s*[-*+]\s+', '  • ', text, flags=re.MULTILINE)
    
    # Remove horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    
    # Clean up tables (simple approach)
    text = re.sub(r'^\|.*\|$', lambda m: m.group(0).replace('|', '  '), text, flags=re.MULTILINE)
    
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text

# Read markdown from artifacts
md_file = '/Users/jp/.gemini/antigravity/brain/b5c0869b-fb33-4631-bf96-6bf87798566d/dossie_imagine_read.md'
with open(md_file, 'r', encoding='utf-8') as f:
    md_content = f.read()

# Convert to text
plain_text = markdown_to_text(md_content)

# Save as TXT in project root
txt_file = Path(__file__).parent / 'dossie_imagine_read.txt'
with open(txt_file, 'w', encoding='utf-8') as f:
    f.write(plain_text)

print(f"✅ Created: {txt_file}")

# Try to create DOCX using docx if available
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Dossiê Técnico - Imagine Read', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content by paragraphs
    for line in plain_text.split('\n'):
        if line.strip():
            if line.startswith('   '):  # H3
                p = doc.add_heading(line.strip(), level=3)
            elif line.startswith('  ') and not line.strip().startswith('•'):  # H2
                p = doc.add_heading(line.strip(), level=2)
            elif not line.startswith(' '):  # H1 or regular text
                if len(line) < 100 and line.isupper():
                    p = doc.add_heading(line.strip(), level=1)
                else:
                    p = doc.add_paragraph(line)
            else:
                p = doc.add_paragraph(line)
    
    # Save DOCX
    docx_file = Path(__file__).parent / 'dossie_imagine_read.docx'
    doc.save(str(docx_file))
    print(f"✅ Created: {docx_file}")
    
except ImportError:
    print("⚠️  python-docx not available, DOCX not generated")
    print("   To install: pip3 install --user python-docx")
